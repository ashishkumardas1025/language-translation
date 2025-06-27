import json
import boto3
from botocore.exceptions import ClientError
import urllib3
from typing import Dict, Any, Optional, Union, List, Tuple
import warnings
import os
import base64
import io
from PIL import Image
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
import camelot
import tabula
import re
from pathlib import Path
import tempfile
import logging

# Configure warnings and disable insecure request warnings
warnings.filterwarnings("ignore", category=UserWarning, message="Unverified HTTPS request")
urllib3.disable_warnings(urllib3.exceptions.InsecureRequestWarning)

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Model configuration
MODEL_ID = "anthropic.claude-3-5-sonnet-20240620-v1:0"

def initialize_bedrock_client():
    """Initialize and return AWS Bedrock client with credentials from environment variables"""
    aws_access_key_id = os.getenv("AWS_ACCESS_KEY_ID")
    aws_secret_access_key = os.getenv("AWS_SECRET_ACCESS_KEY")
    aws_session_token = os.getenv("AWS_SESSION_TOKEN")

    session = boto3.Session(
        aws_access_key_id=aws_access_key_id,
        aws_secret_access_key=aws_secret_access_key,
        aws_session_token=aws_session_token
    )

    bedrock = session.client(service_name='bedrock-runtime', region_name='us-east-1', verify=False)
    return bedrock

def invoke_bedrock_claude(
    prompt: str, 
    system: Optional[str] = None, 
    max_tokens: int = 2048, 
    temperature: float = 0.1,
    images: Optional[List[str]] = None
) -> str:
    """Invoke Claude model through AWS Bedrock with multimodal support"""
    bedrock = initialize_bedrock_client()
    
    # Prepare content array
    content = []
    
    # Add images if provided
    if images:
        for image_base64 in images:
            content.append({
                "type": "image",
                "source": {
                    "type": "base64",
                    "media_type": "image/png",
                    "data": image_base64
                }
            })
    
    # Add text prompt
    content.append({"type": "text", "text": prompt})
    
    request_payload = {
        "anthropic_version": "bedrock-2023-05-31",
        "max_tokens": max_tokens,
        "temperature": temperature,
        "messages": [
            {
                "role": "user",
                "content": content
            }
        ]
    }
    
    if system:
        request_payload["system"] = system
        
    try:
        response = bedrock.invoke_model(
            modelId=MODEL_ID,
            contentType="application/json",
            accept="application/json",
            body=json.dumps(request_payload).encode("utf-8")
        )
        response_body = json.loads(response["body"].read().decode("utf-8"))
        return response_body["content"][0]["text"]
    except ClientError as e:
        logger.error(f"AWS Error: Cannot invoke '{MODEL_ID}'. Reason: {e}")
        raise
    except Exception as e:
        logger.error(f"General Error: {e}")
        raise

class PDFProcessor:
    def __init__(self):
        self.translation_system_prompt = """You are a professional translator specializing in business and financial documents. 
        Translate the provided content from English to Canadian French while maintaining:
        1. Professional terminology and business language
        2. Numerical values and formatting exactly as shown
        3. Technical terms appropriately translated
        4. Document structure and formatting
        5. Cultural adaptation for Canadian French context
        
        Provide only the translated text without explanations or metadata."""

    def extract_pdf_content(self, pdf_path: str) -> Dict[str, Any]:
        """Extract all content from PDF including text, images, and tables"""
        doc = fitz.open(pdf_path)
        content = {
            'pages': [],
            'images': [],
            'tables': [],
            'text_blocks': []
        }
        
        for page_num in range(doc.page_count):
            page = doc.load_page(page_num)
            page_content = {
                'page_number': page_num + 1,
                'text': '',
                'images': [],
                'tables': []
            }
            
            # Extract text
            text = page.get_text()
            page_content['text'] = text
            
            # Extract images
            image_list = page.get_images()
            for img_index, img in enumerate(image_list):
                try:
                    xref = img[0]
                    pix = fitz.Pixmap(doc, xref)
                    if pix.n - pix.alpha < 4:  # GRAY or RGB
                        img_data = pix.tobytes("png")
                        img_base64 = base64.b64encode(img_data).decode()
                        
                        image_info = {
                            'page': page_num + 1,
                            'index': img_index,
                            'base64': img_base64,
                            'width': pix.width,
                            'height': pix.height
                        }
                        page_content['images'].append(image_info)
                        content['images'].append(image_info)
                    pix = None
                except Exception as e:
                    logger.error(f"Error extracting image {img_index} from page {page_num + 1}: {e}")
            
            content['pages'].append(page_content)
            content['text_blocks'].append(text)
        
        # Extract tables using camelot
        try:
            tables = camelot.read_pdf(pdf_path, pages='all', flavor='lattice')
            for i, table in enumerate(tables):
                table_data = {
                    'index': i,
                    'page': table.page,
                    'dataframe': table.df,
                    'csv_string': table.df.to_csv(index=False)
                }
                content['tables'].append(table_data)
        except Exception as e:
            logger.warning(f"Camelot table extraction failed: {e}")
            # Fallback to tabula
            try:
                tables = tabula.read_pdf(pdf_path, pages='all', multiple_tables=True)
                for i, table in enumerate(tables):
                    table_data = {
                        'index': i,
                        'page': i + 1,  # Approximation
                        'dataframe': table,
                        'csv_string': table.to_csv(index=False)
                    }
                    content['tables'].append(table_data)
            except Exception as e2:
                logger.warning(f"Tabula table extraction also failed: {e2}")
        
        doc.close()
        return content

    def translate_text(self, text: str) -> str:
        """Translate text content to Canadian French"""
        if not text.strip():
            return text
            
        try:
            translated = invoke_bedrock_claude(
                prompt=f"Translate this text to Canadian French:\n\n{text}",
                system=self.translation_system_prompt,
                max_tokens=4000
            )
            return translated
        except Exception as e:
            logger.error(f"Translation error: {e}")
            return text  # Return original if translation fails

    def translate_image_content(self, image_base64: str) -> str:
        """Analyze and translate image content using Claude's vision capabilities"""
        try:
            prompt = """Analyze this image and provide a description in Canadian French. 
            If the image contains text, charts, graphs, or diagrams, describe the content and translate any visible text to Canadian French.
            Focus on business/financial content if present."""
            
            description = invoke_bedrock_claude(
                prompt=prompt,
                system=self.translation_system_prompt,
                images=[image_base64],
                max_tokens=1000
            )
            return description
        except Exception as e:
            logger.error(f"Image analysis error: {e}")
            return "Image non analysée"

    def translate_table(self, table_csv: str) -> str:
        """Translate table content to Canadian French"""
        try:
            prompt = f"""Translate this CSV table data to Canadian French while preserving the structure:

{table_csv}

Maintain the CSV format and translate headers and text content appropriately for Canadian French business context."""
            
            translated_csv = invoke_bedrock_claude(
                prompt=prompt,
                system=self.translation_system_prompt,
                max_tokens=2000
            )
            return translated_csv
        except Exception as e:
            logger.error(f"Table translation error: {e}")
            return table_csv

    def create_translated_docx(self, original_content: Dict[str, Any], output_path: str) -> str:
        """Create a translated DOCX document"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Document Traduit - Refuse and Recycling Industry Update', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Process each page
        for page_info in original_content['pages']:
            page_num = page_info['page_number']
            
            # Add page header
            doc.add_heading(f'Page {page_num}', level=1)
            
            # Translate and add text content
            if page_info['text'].strip():
                translated_text = self.translate_text(page_info['text'])
                
                # Split into paragraphs and add to document
                paragraphs = translated_text.split('\n\n')
                for para in paragraphs:
                    if para.strip():
                        doc.add_paragraph(para.strip())
            
            # Add images with descriptions
            for img_info in page_info['images']:
                try:
                    # Add image description
                    img_description = self.translate_image_content(img_info['base64'])
                    doc.add_paragraph(f"Image Description: {img_description}")
                    
                    # Add the actual image
                    img_data = base64.b64decode(img_info['base64'])
                    img_stream = io.BytesIO(img_data)
                    
                    # Calculate appropriate size (max 6 inches width)
                    max_width = Inches(6)
                    aspect_ratio = img_info['height'] / img_info['width']
                    width = min(max_width, Inches(img_info['width'] / 100))
                    height = Inches(width.inches * aspect_ratio)
                    
                    doc.add_picture(img_stream, width=width)
                    doc.add_paragraph("")  # Add spacing
                    
                except Exception as e:
                    logger.error(f"Error adding image to document: {e}")
                    doc.add_paragraph(f"[Image non disponible - Erreur: {str(e)}]")
        
        # Add tables section
        if original_content['tables']:
            doc.add_heading('Tableaux Traduits', level=1)
            
            for table_info in original_content['tables']:
                try:
                    translated_csv = self.translate_table(table_info['csv_string'])
                    
                    # Parse translated CSV and create table
                    lines = translated_csv.strip().split('\n')
                    if lines:
                        # Create table
                        table = doc.add_table(rows=len(lines), cols=len(lines[0].split(',')))
                        table.style = 'Table Grid'
                        
                        for i, line in enumerate(lines):
                            cells = line.split(',')
                            for j, cell in enumerate(cells):
                                if j < len(table.rows[i].cells):
                                    table.rows[i].cells[j].text = cell.strip('"')
                        
                        doc.add_paragraph("")  # Add spacing
                        
                except Exception as e:
                    logger.error(f"Error adding table to document: {e}")
                    doc.add_paragraph(f"[Tableau non disponible - Erreur: {str(e)}]")
        
        # Save document
        doc.save(output_path)
        return output_path

    def process_pdf(self, pdf_path: str, output_dir: str) -> str:
        """Main processing function"""
        logger.info(f"Starting PDF processing: {pdf_path}")
        
        # Extract content
        content = self.extract_pdf_content(pdf_path)
        logger.info(f"Extracted content: {len(content['pages'])} pages, {len(content['images'])} images, {len(content['tables'])} tables")
        
        # Create output path
        pdf_name = Path(pdf_path).stem
        output_path = os.path.join(output_dir, f"{pdf_name}_translated_fr.docx")
        
        # Create translated document
        result_path = self.create_translated_docx(content, output_path)
        logger.info(f"Translation completed: {result_path}")
        
        return result_path

def process_pdf_file(pdf_file_path: str, output_directory: str) -> str:
    """Main function to process PDF file"""
    processor = PDFProcessor()
    return processor.process_pdf(pdf_file_path, output_directory)

if __name__ == "__main__":
    # Test the processor
    pdf_path = "refuse_update_spring_2024_accessible.pdf"
    output_dir = "output"
    
    os.makedirs(output_dir, exist_ok=True)
    
    try:
        result = process_pdf_file(pdf_path, output_dir)
        print(f"Processing completed successfully: {result}")
    except Exception as e:
        print(f"Processing failed: {e}")
